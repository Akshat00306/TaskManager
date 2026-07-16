#!/usr/bin/env python3
"""Generate Summer Training Report DOCX for NextTask (CVM guidelines)."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import copy

OUT = "/home/akshat/Desktop/TaskManager/TaskManager/Summer_Training_Report_NextTask.docx"


def set_run_font(run, size=12, bold=False, italic=False, all_caps=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    if all_caps:
        run.font.all_caps = True


def set_para_format(p, size=12, bold=False, italic=False, align="justify",
                    space_after=Pt(6), space_before=Pt(0), line=1.5,
                    first_indent=None):
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    if line == 1.5:
        pf.line_spacing = 1.5
    elif line == 2:
        pf.line_spacing = 2.0
    elif line == 1:
        pf.line_spacing = 1.0
    else:
        pf.line_spacing = line
    align_map = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    if first_indent is not None:
        pf.first_line_indent = first_indent
    for run in p.runs:
        set_run_font(run, size=size, bold=bold, italic=italic)


def add_para(doc, text, size=12, bold=False, italic=False, align="justify",
             space_after=Pt(8), space_before=Pt(0), line=1.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    set_para_format(p, size=size, bold=bold, italic=italic, align=align,
                    space_after=space_after, space_before=space_before, line=line)
    return p


def add_chapter(doc, text):
    p = add_para(doc, text.upper(), size=16, bold=True, align="center",
                 space_before=Pt(12), space_after=Pt(18), line=1.5)
    return p


def add_section(doc, text):
    return add_para(doc, text.upper(), size=14, bold=True, align="left",
                    space_before=Pt(14), space_after=Pt(10), line=1.5)


def add_subsection(doc, text):
    return add_para(doc, text, size=12, bold=True, align="left",
                    space_before=Pt(10), space_after=Pt(8), line=1.5)


def add_body(doc, text, italic=False):
    return add_para(doc, text, size=12, italic=italic, align="justify",
                    space_after=Pt(10), line=1.5)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    set_run_font(run, size=12)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    return p


def add_page_break(doc):
    doc.add_page_break()


def set_cell_shading(cell, fill="D9D9D9"):
    tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=11, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        # light header shade
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E7E6E6")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15

    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = w

    doc.add_paragraph()
    return table


def add_caption(doc, text, kind="table"):
    # kind: table (above) or figure (below)
    return add_para(doc, text, size=12, bold=False, align="center",
                    space_before=Pt(6), space_after=Pt(10), line=1.5)


def add_header_footer(section, enrollment="[Enrollment No.]",
                      chapter="Summer Training Report",
                      college="[College Name]"):
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_l = hp.add_run(enrollment)
    set_run_font(run_l, size=10)
    # tab to right
    run_tab = hp.add_run("\t")
    set_run_font(run_tab, size=10)
    # set tab stop right
    pPr = hp._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9360")  # ~6.5 inches from left margin area
    tabs.append(tab)
    pPr.append(tabs)
    run_r = hp.add_run(chapter)
    set_run_font(run_r, size=10)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()

    # Three-part footer using a 1x3 table
    ft = footer.add_table(1, 3, Inches(6.5))
    ft.autofit = True
    cells = ft.rows[0].cells
    cells[0].text = ""
    p0 = cells[0].paragraphs[0]
    r0 = p0.add_run("CVM University")
    set_run_font(r0, size=9)
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT

    cells[1].text = ""
    p1 = cells[1].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # PAGE field
    run = p1.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)
    set_run_font(run, size=9)

    cells[2].text = ""
    p2 = cells[2].paragraphs[0]
    r2 = p2.add_run(college)
    set_run_font(r2, size=9)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # remove empty default footer paragraph if present
    if len(footer.paragraphs) > 1:
        pass


def set_black_box_border(section):
    """
    Draw a black rectangular page border (box) inset per CVM margins:
    Left 1.25", Right 1.0", Top 1.0", Bottom 1.0".
    Border is measured from the page edge (w:offsetFrom="page").
    """
    sectPr = section._sectPr
    # Remove existing pgBorders if any
    for child in list(sectPr):
        if child.tag == qn("w:pgBorders"):
            sectPr.remove(child)

    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")

    # Word border "space" is in points (1 inch = 72 pt)
    # Left 1.25" = 90 pt; Right/Top/Bottom 1.0" = 72 pt
    sides = {
        "top": 72,
        "left": 90,
        "bottom": 72,
        "right": 72,
    }
    for side, space_pt in sides.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "24")      # 24 eighths-of-a-point ≈ 1.5 pt line
        el.set(qn("w:space"), str(space_pt))
        el.set(qn("w:color"), "000000")
        pgBorders.append(el)

    sectPr.append(pgBorders)


def configure_section(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    # Content margins sit just inside the black box
    # Box at: L 1.25", R/T/B 1.0" — leave a small gap inside the line
    section.left_margin = Inches(1.35)
    section.right_margin = Inches(1.10)
    section.top_margin = Inches(1.15)
    section.bottom_margin = Inches(1.15)
    set_black_box_border(section)


def build():
    doc = Document()
    section = doc.sections[0]
    configure_section(section)
    add_header_footer(section)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # ========== COVER PAGE ==========
    for _ in range(2):
        add_para(doc, "", size=12, align="center", space_after=Pt(0), line=1)
    add_para(doc, "Summer Training Report", size=18, bold=True, align="center",
             space_after=Pt(24), line=1.5)
    add_para(doc, "[Industry Name]", size=16, bold=True, align="center",
             space_after=Pt(36), line=1.5)
    add_para(doc, "Submitted by", size=12, align="center", space_after=Pt(12), line=1.5)
    add_para(doc, "[NAME OF THE CANDIDATE]", size=14, bold=True, align="center",
             space_after=Pt(6), line=1.5)
    add_para(doc, "[CVMU Enrolment Number]", size=12, align="center",
             space_after=Pt(36), line=1.5)
    add_para(doc, "In partial fulfillment for the award of the degree of",
             size=12, align="center", space_after=Pt(12), line=1.5)
    add_para(doc, "BACHELOR OF TECHNOLOGY", size=14, bold=True, align="center",
             space_after=Pt(6), line=1.5)
    add_para(doc, "in", size=12, align="center", space_after=Pt(6), line=1.5)
    add_para(doc, "[Branch Name]", size=14, bold=True, align="center",
             space_after=Pt(6), line=1.5)
    add_para(doc, "[College Name]", size=12, align="center",
             space_after=Pt(24), line=1.5)
    add_para(doc, "The Charutar Vidya Mandal (CVM) University, Vallabh Vidyanagar - 388120",
             size=12, align="center", space_after=Pt(6), line=1.5)
    add_para(doc, "June, 2026", size=12, bold=True, align="center",
             space_after=Pt(6), line=1.5)

    add_page_break(doc)

    # ========== TITLE PAGE (same) ==========
    add_para(doc, "Summer Training Report", size=18, bold=True, align="center",
             space_before=Pt(48), space_after=Pt(24), line=1.5)
    add_para(doc, "[Industry Name]", size=16, bold=True, align="center",
             space_after=Pt(36), line=1.5)
    add_para(doc, "Submitted by", size=12, align="center", space_after=Pt(12), line=1.5)
    add_para(doc, "[NAME OF THE CANDIDATE]", size=14, bold=True, align="center",
             space_after=Pt(6), line=1.5)
    add_para(doc, "[CVMU Enrolment Number]", size=12, align="center",
             space_after=Pt(36), line=1.5)
    add_para(doc, "In partial fulfillment for the award of the degree of",
             size=12, align="center", space_after=Pt(12), line=1.5)
    add_para(doc, "BACHELOR OF TECHNOLOGY", size=14, bold=True, align="center",
             space_after=Pt(6), line=1.5)
    add_para(doc, "in", size=12, align="center", space_after=Pt(6), line=1.5)
    add_para(doc, "[Branch Name]", size=14, bold=True, align="center",
             space_after=Pt(6), line=1.5)
    add_para(doc, "[College Name]", size=12, align="center",
             space_after=Pt(24), line=1.5)
    add_para(doc, "The Charutar Vidya Mandal (CVM) University, Vallabh Vidyanagar - 388120",
             size=12, align="center", space_after=Pt(6), line=1.5)
    add_para(doc, "June, 2026", size=12, bold=True, align="center", line=1.5)

    add_page_break(doc)

    # ========== COLLEGE CERTIFICATE ==========
    add_para(doc, "[College Name]", size=14, bold=True, align="center", line=2)
    add_para(doc, "[Programme Name]", size=14, bold=True, align="center",
             space_after=Pt(24), line=2)
    add_para(doc, "CERTIFICATE", size=16, bold=True, align="center",
             space_after=Pt(18), line=2)
    add_para(
        doc,
        "This is to certify that [Name of Student] ([EnrollmentNo]) has submitted the "
        "Summer Training report based on internship undergone at [Name of Industry] for a "
        "period of 14 working days from 10/06/2026 to 30/06/2026 in partial fulfillment for "
        "the degree of Bachelor of Technology in [Name of the Branch], [College Name] at "
        "The Charutar Vidya Mandal (CVM) University, Vallabh Vidyanagar during the academic "
        "year 2025 – 26.",
        size=14, align="justify", space_after=Pt(36), line=2,
    )
    for _ in range(3):
        add_para(doc, "", size=14, line=2, space_after=Pt(0))
    add_para(doc, "[Sign]\t\t\t\t\t[Sign]", size=12, align="left", line=1.5)
    add_para(doc, "[Name of Internal Guide]\t\t[Name of Head of the Department]",
             size=12, align="left", line=1.5)
    add_para(doc, "Internal Guide(s)\t\t\tHead of the Department",
             size=12, align="left", line=1.5)

    add_page_break(doc)

    # ========== COMPANY CERTIFICATE ==========
    add_para(doc, "[Industry Letter Head]", size=14, bold=True, align="center", line=1.5)
    add_para(doc, "Date: ____/____/2026", size=12, align="left",
             space_before=Pt(18), line=1.5)
    add_para(doc, "TO WHOM IT MAY CONCERN", size=14, bold=True, align="center",
             space_before=Pt(24), space_after=Pt(18), line=2)
    add_para(
        doc,
        "This is to certify that [Student Name], a student of [Institute Name] has "
        "successfully completed his/her internship in the field of Full-Stack Web "
        "Development / Software Engineering from 10/06/2026 to 30/06/2026 (Total number "
        "of Days: 14) under the guidance of [Industry Mentor].",
        size=12, align="justify", line=1.5,
    )
    add_para(
        doc,
        "His/her internship activities include designing and developing a full-stack task "
        "management system (NextTask) using ASP.NET Core Web API, Angular 18, Entity "
        "Framework Core, SQLite, and JWT-based authentication; implementing role-based "
        "access control, REST APIs, dashboard analytics, and a responsive user interface.",
        size=12, align="justify", line=1.5,
    )
    add_para(
        doc,
        "During the period of her/his internship program with us, he/she had been exposed "
        "to different processes and was found diligent, hardworking, and inquisitive.",
        size=12, align="justify", line=1.5,
    )
    add_para(doc, "We wish him/her every success in his/her life and career.",
             size=12, align="justify", line=1.5)
    add_para(doc, "For [Industry Name]", size=12, align="left",
             space_before=Pt(24), line=1.5)
    add_para(doc, "Authorized Signature with Industry Stamp", size=12,
             align="left", space_before=Pt(36), line=1.5)

    add_page_break(doc)

    # ========== DECLARATION ==========
    add_para(doc, "DECLARATION", size=16, bold=True, align="center",
             space_after=Pt(18), line=2)
    add_para(
        doc,
        "I, [Name of Student] ([EnrollmentNo]), hereby declare that the Summer Training "
        "report submitted in partial fulfillment for the degree of Bachelor of Technology "
        "in [Name of the Branch], [College Name], The Charutar Vidya Mandal (CVM) "
        "University, Vallabh Vidyanagar, is a bonafide record of work carried out by me at "
        "[Industry Name] under the supervision of [External / Internal Guide Name] and that "
        "no part of this report has been directly copied from any students’ reports or taken "
        "from any other source, without providing due reference.",
        size=14, align="justify", line=2,
    )
    for _ in range(4):
        add_para(doc, "", size=14, line=2, space_after=Pt(0))
    add_para(doc, "Name of the Student: [Name of Student]", size=12, align="left", line=1.5)
    add_para(doc, "Sign of Student: ____________________", size=12, align="left", line=1.5)

    add_page_break(doc)

    # ========== ACKNOWLEDGEMENT ==========
    add_para(doc, "ACKNOWLEDGEMENT", size=16, bold=True, align="center",
             space_after=Pt(18), line=2)
    ack = [
        "I take this opportunity to express my sincere gratitude to all those who have helped and guided me during my Summer Training and in the preparation of this report.",
        "I am deeply thankful to [Name of Industry / Organization] for providing me the opportunity to undergo internship and for exposing me to a real software development environment. I am especially grateful to my industry mentor [Industry Mentor Name] for continuous guidance, technical reviews, and encouragement throughout the training period.",
        "I express my heartfelt thanks to my Internal Guide [Internal Guide Name], [College Name], for valuable suggestions and academic support. I am also thankful to the Head of the Department [HoD Name] and all faculty members of the [Department Name] for their cooperation.",
        "I thank The Charutar Vidya Mandal (CVM) University for including Summer Training as part of the curriculum, which helped me bridge classroom learning with industrial practice.",
        "Finally, I thank my family and friends for their constant support and motivation.",
    ]
    for t in ack:
        add_para(doc, t, size=12, align="justify", line=2, space_after=Pt(12))
    add_para(doc, "[Name of Student]", size=12, align="right", space_before=Pt(24), line=1.5)
    add_para(doc, "[Enrollment No.]", size=12, align="right", line=1.5)
    add_para(doc, "[Branch Name]", size=12, align="right", line=1.5)

    add_page_break(doc)

    # ========== ABSTRACT ==========
    add_para(doc, "ABSTRACT", size=16, bold=True, align="center",
             space_after=Pt(18), line=1.5)
    abstract = (
        "This report presents the Summer Training work on NextTask, a full-stack task "
        "management web application. The main objectives were to learn industry tools, "
        "apply software engineering practices, and build a complete system with secure "
        "login, role-based access, REST APIs, and a responsive user interface. "
        "NextTask allows users to create, assign, update, and track tasks with priority, "
        "status, and due date. Admin and User roles are enforced using JWT authentication. "
        "Admins can manage all tasks, while normal users access only their own. A dashboard "
        "shows total, in-progress, completed, and overdue task counts. "
        "The backend uses ASP.NET Core Web API, Entity Framework Core, and SQLite with "
        "BCrypt password hashing. The frontend uses Angular 18 and Angular Material. The "
        "system follows a layered Controller–Service–Repository design. The internship "
        "strengthened practical skills in full-stack development, security, and documentation."
    )
    add_para(doc, abstract, size=14, italic=True, align="justify", line=1.5)
    add_para(doc, "Keywords: Task Management, ASP.NET Core, Angular, JWT, SQLite, Full-Stack Development",
             size=12, italic=True, align="justify", space_before=Pt(12), line=1.5)

    add_page_break(doc)

    # ========== LIST OF FIGURES ==========
    add_para(doc, "LIST OF FIGURES", size=16, bold=True, align="center",
             space_after=Pt(18), line=1.5)
    figs = [
        ("Fig 1.1", "System Context of NextTask"),
        ("Fig 1.2", "High-Level Client–Server Architecture"),
        ("Fig 2.1", "Organizational Structure of the Company"),
        ("Fig 3.1", "Technology Stack of NextTask"),
        ("Fig 3.2", "JWT Authentication Flow"),
        ("Fig 4.1", "Layered Architecture of the Backend API"),
        ("Fig 4.2", "Use Case Diagram of NextTask"),
        ("Fig 4.3", "Entity Relationship Diagram (User–Task)"),
        ("Fig 4.4", "Activity Diagram — Login and Task Management"),
        ("Fig 5.1", "Login Screen of NextTask UI"),
        ("Fig 5.2", "Dashboard Summary Cards"),
        ("Fig 5.3", "Task List with Search and Filters"),
        ("Fig 5.4", "Create / Edit Task Form"),
        ("Fig 5.5", "Sequence Diagram — Create Task API Call"),
    ]
    add_table(doc, ["Figure No.", "Caption", "Page"],
              [(a, b, "") for a, b in figs])

    add_page_break(doc)

    # ========== LIST OF TABLES ==========
    add_para(doc, "LIST OF TABLES", size=16, bold=True, align="center",
             space_after=Pt(18), line=1.5)
    tabs = [
        ("Table 1.1", "Comparison of Manual Tracking vs NextTask"),
        ("Table 3.1", "Backend Technologies and Their Roles"),
        ("Table 3.2", "Frontend Technologies and Their Roles"),
        ("Table 3.3", "Development Tools Used During Internship"),
        ("Table 4.1", "Work Modules Completed During Internship"),
        ("Table 4.2", "User Roles and Permissions"),
        ("Table 4.3", "REST API Endpoints of NextTask"),
        ("Table 5.1", "Attributes of the User Entity"),
        ("Table 5.2", "Attributes of the TaskItem Entity"),
        ("Table 5.3", "Sample Test Cases and Results"),
        ("Table 5.4", "Demo Seed Accounts"),
    ]
    add_table(doc, ["Table No.", "Caption", "Page"],
              [(a, b, "") for a, b in tabs])

    add_page_break(doc)

    # ========== ABBREVIATIONS ==========
    add_para(doc, "LIST OF SYMBOLS, ABBREVIATIONS AND NOMENCLATURE",
             size=16, bold=True, align="center", space_after=Pt(18), line=1.5)
    abbr = [
        ("API", "Application Programming Interface"),
        ("ASP.NET", "Active Server Pages .NET (Microsoft web framework)"),
        ("CLI", "Command Line Interface"),
        ("CORS", "Cross-Origin Resource Sharing"),
        ("CRUD", "Create, Read, Update, Delete"),
        ("CVM", "Charutar Vidya Mandal"),
        ("DI", "Dependency Injection"),
        ("DTO", "Data Transfer Object"),
        ("EF Core", "Entity Framework Core"),
        ("HTTP", "HyperText Transfer Protocol"),
        ("JSON", "JavaScript Object Notation"),
        ("JWT", "JSON Web Token"),
        ("ORM", "Object–Relational Mapping"),
        ("RBAC", "Role-Based Access Control"),
        ("REST", "Representational State Transfer"),
        ("SDK", "Software Development Kit"),
        ("SDLC", "Software Development Life Cycle"),
        ("SPA", "Single Page Application"),
        ("SQL", "Structured Query Language"),
        ("SQLite", "Lightweight embedded relational database"),
        ("UI", "User Interface"),
        ("UX", "User Experience"),
    ]
    add_table(doc, ["Abbreviation", "Description"], abbr)

    add_page_break(doc)

    # ========== TOC ==========
    add_para(doc, "TABLE OF CONTENTS", size=16, bold=True, align="center",
             space_after=Pt(18), line=1.5)
    toc = [
        ("Acknowledgement", "i"),
        ("Abstract", "ii"),
        ("List of Figures", "iii"),
        ("List of Tables", "iv"),
        ("List of Symbols, Abbreviations and Nomenclature", "v"),
        ("Table of Contents", "vi"),
        ("Chapter 1 Introduction", "1"),
        ("Chapter 2 Company Profile", ""),
        ("Chapter 3 Technology Learned", ""),
        ("Chapter 4 Outline of Work Done During Internship", ""),
        ("Chapter 5 Implementation and Results", ""),
        ("Chapter 6 Conclusion and Discussion", ""),
        ("Appendix", ""),
        ("References", ""),
        ("Daily Diary and Attendance Sheet", ""),
    ]
    add_table(doc, ["Content", "Page"], toc)

    add_page_break(doc)

    # ========== CHAPTER 1 ==========
    add_chapter(doc, "CHAPTER 1")
    add_chapter(doc, "INTRODUCTION")

    add_section(doc, "1.1 BACKGROUND")
    add_body(doc,
        "In modern workplaces and academic environments, individuals and small teams need a "
        "reliable way to organize work items, track progress, and meet deadlines. Traditional "
        "methods such as paper notes, shared spreadsheets, and informal chat messages often "
        "fail to provide ownership clarity, status visibility, and secure multi-user access.")
    add_body(doc,
        "Web-based task management systems address these gaps by centralizing task data, "
        "supporting user accounts, and presenting progress through dashboards. During the "
        "Summer Training (10 June 2026 to 30 June 2026), a full-stack application named "
        "NextTask was designed and developed to apply classroom concepts of programming, "
        "databases, networking, and software engineering in a practical industrial-style project.")
    add_body(doc,
        "NextTask is built as a client–server system: an Angular single-page application "
        "communicates with an ASP.NET Core REST API, which persists data in a SQLite "
        "database. Security is implemented using BCrypt password hashing and JWT Bearer "
        "authentication, with role-based rules for Admin and User roles.")
    add_caption(doc, "Fig 1.1   System Context of NextTask")
    add_body(doc, "(Insert diagram: User Browser → Angular UI → ASP.NET Core API → SQLite DB)", italic=True)

    add_section(doc, "1.2 PROBLEM STATEMENT")
    add_body(doc, "Without a dedicated task system, the following problems commonly occur:")
    for b in [
        "Tasks are scattered across multiple tools and messages.",
        "Due dates are missed because overdue items are not highlighted.",
        "There is no clear distinction between personal and administrative oversight.",
        "Progress cannot be summarized quickly for review meetings.",
        "Access control is weak when shared documents are used.",
    ]:
        add_bullet(doc, b)
    add_body(doc,
        "NextTask solves these problems by providing a secure, role-aware, web-based "
        "platform for creating and tracking tasks with searchable lists and dashboard analytics.")
    add_caption(doc, "Table 1.1   Comparison of Manual Tracking vs NextTask")
    add_table(doc,
              ["Aspect", "Manual / Spreadsheet", "NextTask"],
              [
                  ("Central storage", "Weak / fragmented", "Centralized database"),
                  ("User authentication", "Often absent", "JWT + BCrypt"),
                  ("Role control", "Difficult", "Admin / User RBAC"),
                  ("Overdue visibility", "Manual checking", "Automatic overdue count"),
                  ("Dashboard", "Limited", "Live summary cards"),
                  ("Concurrent multi-user use", "Error-prone", "Supported via API"),
              ])

    add_section(doc, "1.3 OBJECTIVES OF THE INTERNSHIP")
    add_body(doc, "The objectives of the Summer Training were:")
    for b in [
        "To understand the software development life cycle in a practical project setting.",
        "To learn and apply ASP.NET Core Web API for building RESTful backend services.",
        "To learn and apply Angular 18 for building a responsive single-page frontend.",
        "To implement secure authentication using JWT and password hashing using BCrypt.",
        "To use Entity Framework Core with SQLite for data modeling and persistence.",
        "To design and implement role-based authorization for Admin and User roles.",
        "To develop CRUD operations, filtering, sorting, and dashboard summary features.",
        "To document the complete system for academic evaluation and future enhancement.",
    ]:
        add_bullet(doc, b)

    add_section(doc, "1.4 SCOPE OF THE PROJECT")
    add_subsection(doc, "1.4.1 In Scope")
    for b in [
        "User registration and login",
        "JWT session management",
        "Task create / read / update / delete",
        "Priority, status, due date, and assignee fields",
        "Admin assignment of tasks to users",
        "Search, filter, sort, and overdue filter",
        "Dashboard summary statistics",
        "Responsive UI with sidenav layout",
        "Seeded demo data for demonstration",
    ]:
        add_bullet(doc, b)
    add_subsection(doc, "1.4.2 Out of Scope (Future Work)")
    for b in [
        "Email / push notifications",
        "File attachments and comments",
        "Multi-tenant organizations / projects",
        "Production database (PostgreSQL / SQL Server)",
        "Refresh-token rotation and advanced security hardening",
    ]:
        add_bullet(doc, b)

    add_section(doc, "1.5 ORGANIZATION OF THE REPORT")
    for b in [
        "Chapter 1 introduces the topic, problem, objectives, and scope.",
        "Chapter 2 presents the company profile and internship environment.",
        "Chapter 3 describes technologies learned during training.",
        "Chapter 4 outlines the work performed and system design.",
        "Chapter 5 explains implementation details and results.",
        "Chapter 6 concludes with discussion, limitations, and future scope.",
        "Appendices, references, and the daily diary follow the chapters.",
    ]:
        add_bullet(doc, b)
    add_caption(doc, "Fig 1.2   High-Level Client–Server Architecture")
    add_body(doc, "(Insert architecture diagram)", italic=True)

    add_page_break(doc)

    # ========== CHAPTER 2 ==========
    add_chapter(doc, "CHAPTER 2")
    add_chapter(doc, "COMPANY PROFILE")

    add_section(doc, "2.1 ABOUT THE ORGANIZATION")
    add_body(doc,
        "[Industry Name] is an organization working in the domain of [software development / "
        "IT services / product engineering]. During the Summer Training period from "
        "10/06/2026 to 30/06/2026, the candidate was associated with the [team / department "
        "name] and contributed to full-stack application development activities related to "
        "the NextTask system.")
    add_body(doc,
        "The organization provided exposure to professional practices such as requirement "
        "discussion, coding standards, code review, testing, and documentation.")

    add_section(doc, "2.2 VISION AND MISSION")
    add_body(doc, "Vision: [Company vision — e.g., to deliver reliable digital products that improve productivity for individuals and teams.]")
    add_body(doc, "Mission: [Company mission — e.g., to build scalable web solutions using modern frameworks while mentoring interns in industry-ready skills.]")

    add_section(doc, "2.3 PRODUCTS / SERVICES")
    add_body(doc, "Typical offerings of the organization include (customize as applicable):")
    for b in [
        "Custom web application development",
        "API and backend engineering",
        "Frontend / UI development",
        "Software maintenance and enhancement",
        "Internship and training mentorship",
    ]:
        add_bullet(doc, b)
    add_body(doc, "The project assigned during internship—NextTask—aligns with the organization’s focus on practical full-stack delivery.")

    add_section(doc, "2.4 WORK ENVIRONMENT AND SUPPORT")
    add_body(doc, "The internship environment included:")
    for b in [
        "Access to development machines and required software / open-source tools",
        "Guidance from an industry mentor for technical doubts and design decisions",
        "Opportunity to work on both backend and frontend layers",
        "Support for debugging, testing, and documenting the solution",
        "Cooperative culture that encouraged questions and iterative improvement",
    ]:
        add_bullet(doc, b)
    add_body(doc,
        "Experience note: The people in the team were cooperative and approachable. Regular "
        "feedback helped improve code quality, naming conventions, and feature completeness. "
        "[Mention if a job / PPO was offered.]")

    add_section(doc, "2.5 LEARNING CULTURE DURING INTERNSHIP")
    add_body(doc, "Learning was encouraged through:")
    for b in [
        "Hands-on coding rather than theory-only sessions",
        "Review of security practices (password hashing, JWT validation)",
        "Understanding of layered architecture and dependency injection",
        "UI/UX considerations such as responsive sidenav behavior",
        "Writing readable APIs and consistent DTO contracts",
    ]:
        add_bullet(doc, b)
    add_caption(doc, "Fig 2.1   Organizational Structure of the Company")
    add_body(doc, "(Insert org chart if available)", italic=True)

    add_page_break(doc)

    # ========== CHAPTER 3 ==========
    add_chapter(doc, "CHAPTER 3")
    add_chapter(doc, "TECHNOLOGY LEARNED")

    add_section(doc, "3.1 OVERVIEW OF TECHNOLOGY STACK")
    add_body(doc, "NextTask uses a modern full-stack combination:")
    for b in [
        "Backend: C#, ASP.NET Core Web API (.NET 10)",
        "Frontend: Angular 18, TypeScript, Angular Material",
        "Database: SQLite via Entity Framework Core",
        "Security: JWT Bearer authentication, BCrypt password hashing",
        "Integration: REST/JSON over HTTP, CORS for local SPA access",
    ]:
        add_bullet(doc, b)
    add_caption(doc, "Fig 3.1   Technology Stack of NextTask")

    add_section(doc, "3.2 BACKEND TECHNOLOGIES")
    add_caption(doc, "Table 3.1   Backend Technologies and Their Roles")
    add_table(doc, ["Technology", "Role in Project"], [
        ("C# / .NET 10", "Core language and runtime for the API"),
        ("ASP.NET Core Web API", "Hosting REST controllers, middleware, DI"),
        ("Entity Framework Core 9", "ORM for Users and Tasks"),
        ("EF Core SQLite provider", "File-based database (nexttask.db)"),
        ("JWT Bearer package", "Token validation and authorization"),
        ("BCrypt.Net-Next", "One-way password hashing"),
        ("OpenAPI", "API documentation endpoint in development"),
    ])
    add_body(doc, "Key learning points:")
    for b in [
        "Configuring services in Program.cs",
        "Creating controllers with route attributes",
        "Separating business logic into services and data access into repositories",
        "Mapping domain models to DTOs for safe API contracts",
        "Global exception handling via custom middleware",
    ]:
        add_bullet(doc, b)

    add_section(doc, "3.3 FRONTEND TECHNOLOGIES")
    add_caption(doc, "Table 3.2   Frontend Technologies and Their Roles")
    add_table(doc, ["Technology", "Role in Project"], [
        ("Angular 18", "SPA framework and component architecture"),
        ("TypeScript", "Typed client-side development"),
        ("Angular Material / CDK", "UI components (forms, sidenav, cards, dialogs)"),
        ("Angular Router", "Navigation, authGuard, guestGuard"),
        ("RxJS", "Observable-based HTTP communication"),
        ("HTTP Interceptor", "Attaching JWT to outgoing API calls"),
        ("Angular CLI / npm", "Project scaffolding, serve, and build"),
    ])

    add_section(doc, "3.4 DATABASE AND SECURITY TECHNOLOGIES")
    add_body(doc,
        "SQLite was chosen for local development because it requires no separate database "
        "server and stores data in a single file (nexttask.db). EF Core creates the schema "
        "using EnsureCreated() and seeds demo users/tasks through DbInitializer.")
    add_body(doc, "Security concepts learned:")
    for b in [
        "Never store plain-text passwords",
        "Issue signed JWTs containing user id, username, and role claims",
        "Validate issuer, audience, signing key, and token lifetime",
        "Enforce authorization both at controller level and service level",
        "Use CORS carefully when SPA and API run on different origins",
    ]:
        add_bullet(doc, b)
    add_caption(doc, "Fig 3.2   JWT Authentication Flow")

    add_section(doc, "3.5 DEVELOPMENT TOOLS")
    add_caption(doc, "Table 3.3   Development Tools Used During Internship")
    add_table(doc, ["Tool", "Purpose"], [
        ("Visual Studio Code / Cursor", "Coding and debugging"),
        (".NET CLI (dotnet)", "Restore, build, and run API"),
        ("Node.js & npm", "Frontend package management"),
        ("Angular CLI (ng)", "Serve and build UI"),
        ("Browser DevTools", "Network and JWT debugging"),
        (".http request file", "Manual API testing"),
        ("Git", "Source version control"),
    ])

    add_page_break(doc)

    # ========== CHAPTER 4 ==========
    add_chapter(doc, "CHAPTER 4")
    add_chapter(doc, "OUTLINE OF WORK DONE DURING INTERNSHIP")

    add_section(doc, "4.1 INTERNSHIP WORK PLAN")
    add_body(doc,
        "The internship was planned for two weeks (14 working days) from 10 June 2026 to "
        "30 June 2026, excluding weekends (Saturdays and Sundays).")
    add_bullet(doc, "Days 1–7 (Week 1): Learning the required technologies and tools")
    add_bullet(doc, "Days 8–14 (Week 2): Building the NextTask project end-to-end")
    add_body(doc, "Internship Duration: 10/06/2026 to 30/06/2026")
    add_body(doc, "Total Working Days: 14 (weekends removed)")
    add_body(doc, "Weekends (no training): 13–14 June, 20–21 June, 27–28 June 2026")

    add_subsection(doc, "4.1.1 Week 1 — Learning Phase (Day 1 to Day 7)")
    add_table(doc, ["Day", "Date", "Weekday", "Focus Area", "Activities / Topics Learned"], [
        ("Day 1", "10/06/2026", "Wednesday", "Orientation & Environment Setup",
         "Understood goals; installed .NET SDK, Node.js, npm, IDE; ran first commands"),
        ("Day 2", "11/06/2026", "Thursday", "C# & ASP.NET Core Basics",
         "C# essentials; Web API structure, middleware, DI, controllers"),
        ("Day 3", "12/06/2026", "Friday", "REST API & HTTP Concepts",
         "REST principles, HTTP methods, status codes, JSON APIs"),
        ("Day 4", "15/06/2026", "Monday", "EF Core & SQLite",
         "ORM, DbContext, entities, relationships, SQLite file DB"),
        ("Day 5", "16/06/2026", "Tuesday", "Authentication & Security",
         "BCrypt hashing; JWT structure, claims, Bearer tokens"),
        ("Day 6", "17/06/2026", "Wednesday", "Angular 18 Fundamentals",
         "Components, TypeScript, data binding, services, HttpClient"),
        ("Day 7", "18/06/2026", "Thursday", "Routing, Material & Design",
         "Router, guards, interceptors, Material; planned layered architecture"),
    ])

    add_subsection(doc, "4.1.2 Week 2 — Development Phase (Day 8 to Day 14)")
    add_table(doc, ["Day", "Date", "Weekday", "Focus Area", "Work Done on NextTask"], [
        ("Day 8", "19/06/2026", "Friday", "Backend Foundation",
         "Created API; SQLite; User/Task models; AppDbContext; repositories"),
        ("Day 9", "22/06/2026", "Monday", "Auth Backend",
         "UserService, TokenService; register/login; BCrypt; JWT generation"),
        ("Day 10", "23/06/2026", "Tuesday", "Task API & RBAC",
         "Task CRUD, DTOs, RBAC, dashboard API, CORS, middleware, seed data"),
        ("Day 11", "24/06/2026", "Wednesday", "Angular Auth UI",
         "nexttask-ui; Login/Register; AuthService; interceptor; guards"),
        ("Day 12", "25/06/2026", "Thursday", "Dashboard & Layout",
         "Shell/sidenav; Dashboard summary cards connected to API"),
        ("Day 13", "26/06/2026", "Friday", "Task UI Features",
         "Task list search/filter/sort/overdue; Create/Edit form; admin assign"),
        ("Day 14", "29/06/2026", "Monday", "Testing & Documentation",
         "E2E testing, bug fixes, demo, README and report notes"),
    ])
    add_body(doc, "(30/06/2026 — Tuesday: within internship end date; used for final review / report submission as applicable.)", italic=True)

    add_subsection(doc, "4.1.3 Summary of Two-Week Plan")
    add_table(doc, ["Week", "Dates", "Nature of Work", "Outcome"], [
        ("Week 1", "10/06/2026 – 18/06/2026 (7 working days)", "Technology learning",
         "Ready with .NET, EF Core, JWT, Angular, Material concepts"),
        ("Week 2", "19/06/2026 – 29/06/2026 (7 working days)", "Project development",
         "Working full-stack NextTask application"),
    ])

    add_caption(doc, "Table 4.1   Work Modules Completed During Internship")
    add_table(doc, ["Module", "Description", "Status"], [
        ("Auth Module", "Register, login, /me, user listing", "Completed"),
        ("Task Module", "Full CRUD + filters", "Completed"),
        ("Dashboard Module", "Summary counts including overdue", "Completed"),
        ("Security Module", "JWT + BCrypt + role checks", "Completed"),
        ("UI Shell", "Responsive sidenav layout", "Completed"),
        ("Seed Data", "Demo admin/user and sample tasks", "Completed"),
    ])

    add_section(doc, "4.2 MODULES DEVELOPED")
    add_subsection(doc, "4.2.1 Authentication Module")
    for b in [
        "POST /api/auth/register — create account",
        "POST /api/auth/login — return JWT and user profile",
        "GET /api/auth/me — validate existing session",
        "GET /api/auth/users — Admin-only user list for assignment dropdown",
    ]:
        add_bullet(doc, b)
    add_subsection(doc, "4.2.2 Task Management Module")
    for b in [
        "List tasks with optional search, status, priority, sortBy, showOverdue",
        "Get task by id; Create / update / delete task",
        "Admin can assign tasks to other users; users are self-assigned",
    ]:
        add_bullet(doc, b)
    add_subsection(doc, "4.2.3 Dashboard Module")
    add_body(doc, "Returns counts for: Total tasks; Todo / InProgress / Completed; Overdue tasks.")
    add_subsection(doc, "4.2.4 Frontend Feature Modules")
    for b in [
        "Login and Register pages with guest guard",
        "Authenticated shell with navigation",
        "Dashboard cards",
        "Task list and create/edit form",
    ]:
        add_bullet(doc, b)

    add_section(doc, "4.3 SYSTEM ARCHITECTURE")
    add_body(doc, "The backend follows a clean layered approach: Controllers → Services → Repositories → EF Core / SQLite.")
    add_caption(doc, "Fig 4.1   Layered Architecture of the Backend API")
    add_body(doc, "Benefits observed during implementation:")
    for b in [
        "Controllers remain thin and HTTP-focused",
        "Business rules (especially authorization) stay in services",
        "Repositories encapsulate queries",
        "Easier testing and future replacement of storage",
    ]:
        add_bullet(doc, b)
    add_caption(doc, "Fig 4.2   Use Case Diagram of NextTask")
    add_caption(doc, "Fig 4.3   Entity Relationship Diagram (User–Task)")
    add_caption(doc, "Fig 4.4   Activity Diagram — Login and Task Management")

    add_section(doc, "4.4 ROLE-BASED ACCESS DESIGN")
    add_caption(doc, "Table 4.2   User Roles and Permissions")
    add_table(doc, ["Capability", "Admin", "User"], [
        ("View all tasks", "Yes", "No (own only)"),
        ("Create task", "Yes", "Yes (self-assigned)"),
        ("Assign to another user", "Yes", "No"),
        ("Edit / delete any task", "Yes", "Own only"),
        ("List all users", "Yes", "No"),
        ("Dashboard scope", "All tasks", "Own tasks"),
    ])

    add_section(doc, "4.5 API DESIGN SUMMARY")
    add_body(doc, "Base URL (development): http://localhost:5292/api")
    add_body(doc, "Frontend: http://localhost:4200")
    add_caption(doc, "Table 4.3   REST API Endpoints of NextTask")
    add_table(doc, ["Method", "Endpoint", "Auth", "Description"], [
        ("POST", "/auth/register", "No", "Register user"),
        ("POST", "/auth/login", "No", "Login and receive JWT"),
        ("GET", "/auth/me", "Yes", "Current user / session check"),
        ("GET", "/auth/users", "Admin", "List users"),
        ("GET", "/tasks", "Yes", "List/filter tasks"),
        ("GET", "/tasks/{id}", "Yes", "Get task"),
        ("POST", "/tasks", "Yes", "Create task"),
        ("PUT", "/tasks/{id}", "Yes", "Update task"),
        ("DELETE", "/tasks/{id}", "Yes", "Delete task"),
        ("GET", "/tasks/dashboard", "Yes", "Dashboard summary"),
    ])

    add_page_break(doc)

    # ========== CHAPTER 5 ==========
    add_chapter(doc, "CHAPTER 5")
    add_chapter(doc, "IMPLEMENTATION AND RESULTS")

    add_section(doc, "5.1 SYSTEM REQUIREMENTS")
    add_body(doc, "Hardware (minimum): Dual-core CPU, 8 GB RAM, 2 GB free disk space.")
    add_body(doc, "Software: .NET 10 SDK; Node.js 18+ and npm; Modern browser; OS: Windows / Linux / macOS.")
    add_body(doc, "How to run:")
    add_body(doc, "Backend: cd TaskManager → dotnet restore → dotnet run --launch-profile http", italic=True)
    add_body(doc, "Frontend: cd nexttask-ui → npm install → npm start", italic=True)

    add_section(doc, "5.2 DATABASE DESIGN")
    add_subsection(doc, "5.2.1 User Entity")
    add_caption(doc, "Table 5.1   Attributes of the User Entity")
    add_table(doc, ["Field", "Type", "Constraints / Notes"], [
        ("Id", "int", "Primary key"),
        ("Username", "string", "Required, unique index"),
        ("PasswordHash", "string", "BCrypt hash"),
        ("Role", "enum", "Admin / User"),
        ("CreatedDate", "DateTime", "UTC timestamp"),
    ])
    add_subsection(doc, "5.2.2 TaskItem Entity")
    add_caption(doc, "Table 5.2   Attributes of the TaskItem Entity")
    add_table(doc, ["Field", "Type", "Constraints / Notes"], [
        ("Id", "int", "Primary key"),
        ("Title", "string", "Required, 3–100 characters"),
        ("Description", "string?", "Optional"),
        ("Priority", "enum", "Low / Medium / High"),
        ("Status", "enum", "Todo / InProgress / Completed"),
        ("DueDate", "DateTime?", "Optional"),
        ("CreatedDate", "DateTime", "Set on create"),
        ("UpdatedDate", "DateTime", "Updated on modify"),
        ("AssignedUserId", "int?", "FK to User (SetNull on delete)"),
    ])
    add_body(doc, "Relationship: one User may have many assigned TaskItems.")

    add_section(doc, "5.3 BACKEND IMPLEMENTATION")
    add_body(doc, "Important implementation highlights:")
    for b in [
        "Dependency Injection registers repositories and services as scoped dependencies.",
        "JWT configuration validates issuer (NextTaskAPI), audience (NextTaskUI), signing key, and lifetime.",
        "ExceptionMiddleware converts exceptions into consistent JSON responses (401, 404, 400, 500).",
        "DbInitializer ensures database creation and seeds admin/john accounts with sample tasks.",
        "TaskService enforces that non-admin users cannot view, edit, or delete tasks belonging to others.",
    ]:
        add_bullet(doc, b)

    add_section(doc, "5.4 FRONTEND IMPLEMENTATION")
    for b in [
        "AuthService stores JWT and user session; validates expiry.",
        "authGuard blocks unauthenticated access to dashboard/tasks.",
        "guestGuard redirects already-logged-in users away from login/register.",
        "Auth interceptor attaches Authorization: Bearer <token> automatically.",
        "DashboardComponent loads summary cards from /tasks/dashboard.",
        "TaskListComponent supports search, status/priority filters, sorting, and overdue filter.",
        "TaskFormComponent supports create/edit; Admin can select assignee.",
    ]:
        add_bullet(doc, b)

    add_section(doc, "5.5 TESTING AND RESULTS")
    add_caption(doc, "Table 5.3   Sample Test Cases and Results")
    add_table(doc, ["Test Case", "Steps", "Expected Result", "Actual Result"], [
        ("TC01 Valid Admin Login", "Login as admin", "JWT returned; dashboard opens", "Pass"),
        ("TC02 Invalid Password", "Wrong password", "Error; no session", "Pass"),
        ("TC03 User Isolation", "Login as john; list tasks", "Only john’s tasks visible", "Pass"),
        ("TC04 Admin Sees All", "Login as admin; list tasks", "All tasks visible", "Pass"),
        ("TC05 Create Task", "Submit valid form", "Task appears in list", "Pass"),
        ("TC06 Overdue Filter", "Enable showOverdue", "Only overdue tasks listed", "Pass"),
        ("TC07 Unauthorized Edit", "User edits another’s task", "401 / unauthorized", "Pass"),
        ("TC08 Session Restore", "Refresh with valid token", "Remains logged in via /me", "Pass"),
        ("TC09 Delete Task", "Delete own task", "Task removed (204)", "Pass"),
        ("TC10 Admin Assign", "Admin assigns task to john", "Assignee updated", "Pass"),
    ])
    add_caption(doc, "Table 5.4   Demo Seed Accounts")
    add_table(doc, ["Username", "Password", "Role"], [
        ("admin", "Admin@123", "Admin"),
        ("john", "User@123", "User"),
    ])

    add_section(doc, "5.6 SCREENSHOTS / OUTPUT")
    add_body(doc, "(Paste actual screenshots in the Word file under the captions below.)", italic=True)
    for cap in [
        "Fig 5.1   Login Screen of NextTask UI",
        "Fig 5.2   Dashboard Summary Cards",
        "Fig 5.3   Task List with Search and Filters",
        "Fig 5.4   Create / Edit Task Form",
        "Fig 5.5   Sequence Diagram — Create Task API Call",
    ]:
        add_caption(doc, cap)
        add_body(doc, "[Insert screenshot / diagram here]", italic=True)
    add_body(doc,
        "Observed results: The integrated system successfully supports secure multi-user "
        "task management with clear role separation and usable dashboard analytics. API "
        "responses are JSON-based and suitable for SPA consumption. The SQLite database "
        "initializes automatically, which simplifies demonstration and evaluation.")

    add_page_break(doc)

    # ========== CHAPTER 6 ==========
    add_chapter(doc, "CHAPTER 6")
    add_chapter(doc, "CONCLUSION AND DISCUSSION")

    add_section(doc, "6.1 CONCLUSION")
    add_body(doc,
        "The Summer Training successfully achieved its learning and delivery objectives. A "
        "working full-stack product, NextTask, was implemented using ASP.NET Core, Angular 18, "
        "EF Core, SQLite, JWT, and BCrypt. The system demonstrates practical understanding of "
        "REST APIs, authentication, authorization, ORM-based data access, and responsive UI "
        "development.")
    add_body(doc,
        "The internship bridged academic knowledge with industrial practices such as layered "
        "architecture, secure coding, iterative debugging, and documentation. The completed "
        "application can be demonstrated end-to-end using seeded accounts and is suitable as "
        "a portfolio project for further enhancement.")

    add_section(doc, "6.2 DISCUSSION")
    add_body(doc,
        "Developing NextTask highlighted that security and authorization must be designed "
        "early, not added at the end. Dual enforcement of roles at API and service layers "
        "prevented data leaks between users. Using DTOs avoided over-exposing domain entities. "
        "On the frontend, guards and interceptors significantly improved session handling "
        "compared to ad-hoc token checks in each component.")
    add_body(doc,
        "SQLite was excellent for training and demos, though production deployments would "
        "typically migrate to a server-based RDBMS. Angular Material accelerated UI "
        "development while keeping the interface consistent.")

    add_section(doc, "6.3 LIMITATIONS")
    for b in [
        "SQLite is local-file based and not ideal for high-concurrency production hosting.",
        "JWT secret is stored in configuration suitable for development only.",
        "No email verification or password-reset workflow.",
        "No real-time updates (SignalR / WebSockets).",
        "Limited automated unit/integration test coverage.",
        "No multi-project or team workspace model yet.",
    ]:
        add_bullet(doc, b)

    add_section(doc, "6.4 FUTURE SCOPE")
    for b in [
        "Add email reminders for overdue tasks.",
        "Introduce comments, attachments, and activity history.",
        "Support projects/teams with finer permissions.",
        "Migrate to PostgreSQL/SQL Server and containerize with Docker.",
        "Implement refresh tokens and stronger secret management.",
        "Expand automated testing (xUnit / Jasmine-Karma) and CI/CD pipelines.",
        "Add analytics charts for productivity trends.",
    ]:
        add_bullet(doc, b)

    add_page_break(doc)

    # ========== APPENDIX ==========
    add_chapter(doc, "APPENDIX")
    add_section(doc, "APPENDIX 1  SAMPLE API REQUEST — LOGIN")
    add_body(doc, "POST http://localhost:5292/api/auth/login", italic=True)
    add_body(doc, 'Body: { "username": "admin", "password": "Admin@123" }', italic=True)

    add_section(doc, "APPENDIX 2  SAMPLE API REQUEST — CREATE TASK")
    add_body(doc, "POST http://localhost:5292/api/tasks", italic=True)
    add_body(doc, "Authorization: Bearer <token>", italic=True)
    add_body(doc,
        'Body: { "title": "Review pull requests", "description": "Go through open PRs", '
        '"priority": "High", "status": "Todo", "dueDate": "2026-07-20T00:00:00Z" }',
        italic=True)

    add_section(doc, "APPENDIX 3  PROJECT FOLDER SNAPSHOT")
    add_bullet(doc, "TaskManager/  → ASP.NET Core API (NextTaskAPI)")
    add_bullet(doc, "nexttask-ui/  → Angular 18 frontend")

    add_page_break(doc)

    # ========== REFERENCES ==========
    add_para(doc, "REFERENCES", size=16, bold=True, align="center",
             space_after=Pt(18), line=1)
    refs = [
        "1. Angular Team (2024) ‘Angular Documentation’, Available at: https://angular.dev/ (Accessed: 30 June 2026).",
        "2. Fielding, R.T. (2000) ‘Architectural Styles and the Design of Network-based Software Architectures’, Doctoral Dissertation, University of California, Irvine. Available at: https://www.ics.uci.edu/~fielding/pubs/dissertation/top.htm (Accessed: 30 June 2026).",
        "3. IETF (2015) ‘JSON Web Token (JWT)’, RFC 7519, Available at: https://datatracker.ietf.org/doc/html/rfc7519 (Accessed: 30 June 2026).",
        "4. Microsoft (2024) ‘ASP.NET Core Documentation’, Available at: https://learn.microsoft.com/aspnet/core/ (Accessed: 30 June 2026).",
        "5. Microsoft (2024) ‘Entity Framework Core Documentation’, Available at: https://learn.microsoft.com/ef/core/ (Accessed: 30 June 2026).",
        "6. Microsoft (2024) ‘JWT Bearer Authentication in ASP.NET Core’, Available at: https://learn.microsoft.com/aspnet/core/security/authentication/jwt-authn (Accessed: 30 June 2026).",
        "7. Mozilla Developer Network (2024) ‘HTTP Overview’, Available at: https://developer.mozilla.org/en-US/docs/Web/HTTP/Overview (Accessed: 30 June 2026).",
        "8. OpenJS Foundation (2024) ‘Node.js Documentation’, Available at: https://nodejs.org/en/docs (Accessed: 30 June 2026).",
        "9. OWASP Foundation (2024) ‘Authentication Cheat Sheet’, Available at: https://cheatsheetseries.owasp.org/cheatsheets/Authentication_Cheat_Sheet.html (Accessed: 30 June 2026).",
        "10. SQLite Development Team (2024) ‘SQLite Documentation’, Available at: https://www.sqlite.org/docs.html (Accessed: 30 June 2026).",
    ]
    for r in refs:
        add_para(doc, r, size=12, align="left", line=1, space_after=Pt(8))

    add_page_break(doc)

    # ========== DAILY DIARY ==========
    add_chapter(doc, "DAILY DIARY AND ATTENDANCE SHEET")
    add_body(doc, "Period: 10/06/2026 to 30/06/2026")
    add_body(doc, "Working days: 14 (Saturdays & Sundays excluded)")
    add_body(doc, "Week 1 = Learning | Week 2 = Project Building")

    add_subsection(doc, "Week 1 — Technology Learning (Day 1 to Day 7)")
    add_table(doc, ["Day / Date", "Work Done", "Learning / Remarks", "Mentor Initials"], [
        ("Day 1 — Wed, 10/06/2026", "Orientation; installed .NET SDK, Node.js, npm, IDE; reviewed requirements", "Tools and environment setup", ""),
        ("Day 2 — Thu, 11/06/2026", "Studied C# basics and ASP.NET Core Web API", "Backend framework fundamentals", ""),
        ("Day 3 — Fri, 12/06/2026", "Learned REST, HTTP methods, status codes, JSON APIs", "API communication concepts", ""),
        ("Weekend 13–14/06/2026", "No training", "Weekend", ""),
        ("Day 4 — Mon, 15/06/2026", "Studied EF Core, entities, DbContext, SQLite", "ORM and database basics", ""),
        ("Day 5 — Tue, 16/06/2026", "Learned BCrypt and JWT authentication flow", "Security fundamentals", ""),
        ("Day 6 — Wed, 17/06/2026", "Studied Angular 18, TypeScript, services, HttpClient", "Frontend framework basics", ""),
        ("Day 7 — Thu, 18/06/2026", "Learned Router, guards, interceptors, Material; planned architecture", "UI routing + project design ready", ""),
    ])

    add_subsection(doc, "Week 2 — Project Building (Day 8 to Day 14)")
    add_table(doc, ["Day / Date", "Work Done", "Learning / Remarks", "Mentor Initials"], [
        ("Day 8 — Fri, 19/06/2026", "Created API project; SQLite; User/Task models; DbContext", "Backend foundation started", ""),
        ("Weekend 20–21/06/2026", "No training", "Weekend", ""),
        ("Day 9 — Mon, 22/06/2026", "Implemented register/login, BCrypt, TokenService, JWT", "Auth module completed", ""),
        ("Day 10 — Tue, 23/06/2026", "Task CRUD + RBAC + dashboard API; CORS; middleware; seed data", "Core backend completed", ""),
        ("Day 11 — Wed, 24/06/2026", "Angular app; Login/Register UI; AuthService; interceptor; guards", "Frontend auth completed", ""),
        ("Day 12 — Thu, 25/06/2026", "App shell/sidenav and Dashboard summary cards", "Dashboard UI completed", ""),
        ("Day 13 — Fri, 26/06/2026", "Task List (search/filter/sort/overdue) and Create/Edit form", "Task UI completed", ""),
        ("Weekend 27–28/06/2026", "No training", "Weekend", ""),
        ("Day 14 — Mon, 29/06/2026", "E2E testing, bug fixes, demo, documentation/report notes", "Project finalized", ""),
        ("30/06/2026", "Internship end date; final review / submission", "As applicable", ""),
    ])

    add_subsection(doc, "Attendance Sheet")
    add_table(doc, ["Week", "From", "To", "Present Days", "Absent Days", "Mentor Sign"], [
        ("Week 1 (Learning)", "10/06/2026", "18/06/2026", "7", "0", ""),
        ("Week 2 (Building)", "19/06/2026", "29/06/2026", "7", "0", ""),
        ("Total", "10/06/2026", "30/06/2026", "14", "0", ""),
    ])
    add_body(doc, "Declaration: I confirm that the above diary and attendance record are true to the best of my knowledge.")
    add_body(doc, "Student Sign: ______________     Mentor Sign: ______________     Date: __________")

    add_page_break(doc)
    add_para(doc, "NOTE TO STUDENT", size=14, bold=True, align="center", space_after=Pt(12))
    for b in [
        "Replace all [SQUARE BRACKETS] with your real name, enrollment number, college, branch, industry, and guide names.",
        "Insert screenshots under Fig 5.1–5.4 and diagrams where captions appear.",
        "Get College Certificate, Company Certificate, Declaration, and Diary signed.",
        "Export PDF as <EnrollmentNo>_2026.pdf and spiral-bind hard copies.",
        "Update header Top-Left with your Enrollment Number and Bottom-Right with College Name in Word.",
    ]:
        add_bullet(doc, b)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
